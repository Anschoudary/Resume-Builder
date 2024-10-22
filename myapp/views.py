from django.shortcuts import render, redirect
from .models import Resume, User, Education, Experience, Project, Skill, Tool
from django.http import HttpResponse
import os
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from docx import Document
import re

# Create your views here.

def home(request):
    return render(request, 'index.html')

def logout(request):
    # Clear the session
    request.session.flush()
    # Redirect to home or login page
    return redirect('login')

def login(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        password = request.POST.get('password')

        # Check if the user already exists
        user = User.objects.filter(name=name).first()

        if user:
            # If user exists, check if the password matches
            if user.password == password:
                request.session['user_id'] = user.id  # Store user id in session
                return redirect('home')
            else:
                error = 'Invalid credentials. Please try again.'
                return render(request, 'login.html', {'error': error})
        else:
            # If user doesn't exist, create a new user
            user = User.objects.create(name=name, password=password)
            request.session['user_id'] = user.id  # Store user id in session
            return redirect('home')

    return render(request, 'login.html')

def build_resume(request):
    user_id = request.session.get('user_id')
    if not user_id:
        return redirect('login')  # Redirect to login if not logged in

    user = User.objects.get(id=user_id)
    resume = Resume.objects.filter(user=user).first()

    # Retrieve the selected template either from session or URL parameters
    selected_template = request.GET.get('template')
    if selected_template:
        request.session['selected_template'] = selected_template
    else:
        selected_template = request.session.get('selected_template', 'default_template')  # Fallback to default if not set

    if request.method == 'POST':
        # Basic resume data
        data = {
            'name': request.POST.get('name'),
            'email': request.POST.get('email'),
            'phone': request.POST.get('phone'),
            'profile': request.POST.get('profile', ''),
            'about': request.POST.get('about', ''),
            'user': user,
        }

        if resume:
            # Update existing resume
            for key, value in data.items():
                setattr(resume, key, value)
            resume.save()
        else:
            # Create new resume
            resume = Resume.objects.create(**data)

        # Clear existing related entries before saving new ones
        resume.educations.all().delete()
        resume.experiences.all().delete()
        resume.projects.all().delete()
        resume.skills.all().delete()
        resume.tools.all().delete()

        # Save education entries
        degrees = request.POST.getlist('degree[]')
        colleges = request.POST.getlist('college[]')
        degree_years = request.POST.getlist('degree_year[]')

        for degree, college, degree_year in zip(degrees, colleges, degree_years):
            if degree and college and degree_year:
                Education.objects.create(resume=resume, degree=degree, college=college, degree_year=degree_year)

        # Save experience entries
        organizations = request.POST.getlist('company[]')
        designations = request.POST.getlist('designation[]')
        exp_years = request.POST.getlist('exp_year[]')

        for organization, designation, exp_year in zip(organizations, designations, exp_years):
            if organization and designation and exp_year:
                Experience.objects.create(resume=resume, organization=organization, designation=designation, exp_year=exp_year)

        # Save project entries
        projects = request.POST.getlist('project[]')
        descriptions = request.POST.getlist('desc[]')

        for project, description in zip(projects, descriptions):
            if project:
                Project.objects.create(resume=resume, project_name=project, description=description)

        # Save skill entries
        skills = request.POST.getlist('skill[]')

        for skill in skills:
            if skill:
                Skill.objects.create(resume=resume, skill_name=skill)

        # Save tool entries
        tools = request.POST.getlist('tool[]')

        for tool in tools:
            if tool:
                Tool.objects.create(resume=resume, tool_name=tool)

        # Redirect to user_resume after saving
        return redirect('user_resume', username=user.name)

    # If the user already has a resume, pre-fill the form fields with existing data
    if resume:
        educations = resume.educations.all()
        experiences = resume.experiences.all()
        projects = resume.projects.all()
        skills = resume.skills.all()
        tools = resume.tools.all()

        data = {
            'name': resume.name,
            'email': resume.email,
            'phone': resume.phone,
            'profile': resume.profile,
            'about': resume.about,
            'educations': educations,
            'experiences': experiences,
            'projects': projects,
            'skills': skills,
            'tools': tools,
            'selected_template': selected_template  # Pass the selected template to the template
        }

        return render(request, 'resume.html', {'data': data})

    else:
        data = {
            'name': '',
            'email': '',
            'phone': '',
            'profile': '',
            'about': '',
            'educations': [],
            'experiences': [],
            'projects': [],
            'skills': [],
            'tools': [],
            'selected_template': selected_template  # Pass the selected template to the template
        }

    return render(request, 'resume.html', {'data': data})
    
def demo_templates(request):
    user_id = request.session.get('user_id')
    if not user_id:
        return redirect('login')

    if request.method == 'POST':
        # Get the selected template from the form POST data
        selected_template = request.POST.get('template_choice')
        
        if selected_template:
            # Save the selected template in the session
            request.session['selected_template'] = selected_template
            
            # Redirect to build_resume view after selecting the template
            return redirect('build_resume')
    else:
        # If no template is selected, set a default template
        if 'selected_template' not in request.session:
            request.session['selected_template'] = 'default_template'  # Replace 'default_template' with your actual default template name

    return render(request, 'demo_templates.html')

def user_resume(request, username):
    user = User.objects.filter(name=username).first()
    if not user:
        return redirect('login')  # Redirect to login if the user is not found

    resume = Resume.objects.filter(user=user).first()
    selected_template = request.session.get('selected_template')  # Ensure a default template is always set

    user_data = {
        'name': resume.name,
        'email': resume.email,
        'phone': resume.phone,
        'profile': resume.profile,
        'about': resume.about,
        'educations': ', '.join([f"{edu.degree} from {edu.college} ({edu.degree_year})" for edu in resume.educations.all()]),
        'experiences': ', '.join([f"{exp.designation} at {exp.organization} ({exp.exp_year})" for exp in resume.experiences.all()]),
        'projects': ', '.join([f"{proj.project_name}: {proj.description}" for proj in resume.projects.all()]),
        'skills': ', '.join([skill.skill_name for skill in resume.skills.all()]),
        'tools': ', '.join([tool.tool_name for tool in resume.tools.all()]),
    }

    replace_placeholders_in_docx(
        f'/home/m-ans/Documents/Python/Resume_builder/myapp/static/templates/classic_resume.docx',
        f'/home/m-ans/Documents/Python/Resume_builder/myapp/static/templates/Resume.docx',
        user_data
    )

    if resume:
        data = {
            'name': resume.name,
            'email': resume.email,
            'phone': resume.phone,
            'profile': resume.profile,
            'about': resume.about,
            'educations': resume.educations.all(),
            'experiences': resume.experiences.all(),
            'projects': resume.projects.all(),
            'skills': resume.skills.all(),
            'tools': resume.tools.all(),
        }
        return render(request, f'default_template.html', {'data': data})

    return redirect('build_resume')

def download_template(request, template_name):
    user_id = request.session.get('user_id')
    if not user_id:
        return redirect('login')

    user = User.objects.get(id=user_id)
    resume = Resume.objects.filter(user=user).first()
    selected_template = template_name

    if not selected_template:
        return HttpResponse("No template selected", status=400)

    if resume:
        data = {
            'name': resume.name,
            'email': resume.email,
            'phone': resume.phone,
            'profile': resume.profile,
            'about': resume.about,
            'educations': resume.educations.all(),
            'experiences': resume.experiences.all(),
            'projects': resume.projects.all(),
            'skills': resume.skills.all(),
            'tools': resume.tools.all(),
        }
        file_path = '/home/m-ans/Documents/Python/Resume_builder/myapp/static/templates/Resume.pdf'
        if os.path.exists(file_path):
            with open(file_path, 'rb') as pdf_file:
                response = HttpResponse(pdf_file.read(), content_type='application/pdf')
                response['Content-Disposition'] = 'attachment; filename="Resume.pdf"'
            return response
        else:
            return HttpResponse("The requested file was not found", status=404)

    return redirect('build_resume')


def replace_placeholders_in_docx(template_path, output_path, user_data):
    """
    Replaces placeholders in the .docx template with user data.

    :param template_path: Path to the .docx template file.
    :param output_path: Path to save the updated .docx file.
    :param user_data: Dictionary containing user data to replace in the template.
                      Example: {'name': 'John Doe', 'email': 'john@example.com', ...}
    :return: None
    """
    # Load the .docx template
    doc = Document(template_path)

    
    placeholder_pattern = re.compile(r"\{\{(\w+)\}\}")

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        if placeholder_pattern.search(paragraph.text):
            for match in placeholder_pattern.findall(paragraph.text):
                if match in user_data:
                    placeholder = f"{{{{{match}}}}}"  # Construct the placeholder like {{name}}
                    paragraph.text = paragraph.text.replace(placeholder, user_data[match])

    # Replace placeholders in tables (if there are any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder_pattern.search(cell.text):
                    for match in placeholder_pattern.findall(cell.text):
                        if match in user_data:
                            placeholder = f"{{{{{match}}}}}"
                            cell.text = cell.text.replace(placeholder, user_data[match])

    # Save the updated document to the output path
    doc.save(output_path)

    print(f"Resume generated: {output_path}")
